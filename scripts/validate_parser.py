#!/usr/bin/env python3
"""
Validator + reference implementation for the team-document parsers.

Run from repo root:
    python3 tdf-pool-v2/scripts/validate_parser.py

Reads the historical files in the parent folder and emits a per-year JSON
summary so we can eyeball that the TypeScript parser (lib/parsers/*.ts)
will agree with what's in those files.

Document structure (validated across 2020/2021/2022/2024/2025):

  Header text (rules + scoring tables)
  ... then per team, in this order:
    1) "Player's Team Name" paragraph             ← may be missing in messy docs
    2) Rider table (col 0 = rider name)
    3) "Reserve's: 1) X 2) Y 3) Z" paragraph      ← may be glommed onto next team's header

Edge cases handled:
  - Some teams' header is concatenated to the prior reserves paragraph
    (e.g. 2024 has "...3) GallEelco is still trying"). We extract the trailing
    header and backfill it onto the most recent headerless team.
  - Some teams have no detectable header at all → recorded as "Unknown_N"
    and surfaced in `unresolved`. The /admin/upload UI lets Sofia rename
    them before confirming the import.
"""
from __future__ import annotations

import csv
import json
import re
import sys
from pathlib import Path
from typing import List, Dict, Optional, Tuple

try:
    import docx  # python-docx
except ImportError:
    print("Need python-docx: pip install python-docx --break-system-packages")
    sys.exit(1)


REPO = Path(__file__).resolve().parents[1]
INPUTS_DIR = REPO / "historical-inputs"

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

YEAR_RE = re.compile(r"\bTour\s+(\d{4})\b")
# Possessive header: "Quinten's Let's Win again". Greedy first capture so
# multi-word player names like "Bas Otto" still work.
HEADER_RE = re.compile(r"^([A-Z][A-Za-z][A-Za-z .\-]*?)\s*[’']\s*s\s+(.+)$")
RESERVES_LINE_RE = re.compile(r"^\s*reserve", re.IGNORECASE)
# A reserve entry: "1) Mas", "1)Mas", "1.Mas"
RESERVE_ITEM_RE = re.compile(r"\d+\s*[\)\.]\s*([A-Za-z][A-Za-z .\-’']*?)(?=\s*(?:\d+\s*[\)\.]|$))")
SKIP_HEADERS = ("renner/etappe", "dagtotaal", "cumulatief", "totaal")
HEADER_NOISE = (
    "dear teamleaders", "the rules", "the scoring", "etappe placing",
    "have fun", "remember at the start", "final placing",
)
# These tokens routinely appear inside reserves paragraphs as part of names —
# don't get confused into thinking they're "Player's Team".
RESERVE_INDICATOR = ("reserve",)


def find_year(text: str) -> Optional[int]:
    m = YEAR_RE.search(text)
    return int(m.group(1)) if m else None


def is_noise(text: str) -> bool:
    t = text.lower()
    return any(n in t for n in HEADER_NOISE)


def extract_header(text: str) -> Optional[Tuple[str, str]]:
    """Return (player, team_name) if `text` is a clean possessive header."""
    line = text.strip()
    if not line or len(line) > 120 or is_noise(line):
        return None
    if any(t in line.lower() for t in RESERVE_INDICATOR):
        return None
    m = HEADER_RE.match(line)
    if not m:
        return None
    player, team = m.group(1).strip(), m.group(2).strip()
    if not player or len(player) > 30 or len(team) > 100:
        return None
    return player, team


def parse_reserves_paragraph(text: str) -> Tuple[List[str], Optional[Tuple[str, str]]]:
    """
    Split a "Reserve's: ..." paragraph into:
      - the list of reserve names (cleaned)
      - any trailing team-header that got glommed on (or None)

    Examples:
      "Reserve's: 1) Mas  2) Onley  3) S Yates"
        → (["Mas","Onley","S Yates"], None)
      "Reserve's:  1) A Yates 2) Rodriguez 3) Gall  Eelco's tour ploeg"
        → (["A Yates","Rodriguez","Gall"], ("Eelco","tour ploeg"))
      "Reserve's: 1) Carapaz 2) Mohoric 3) GallEelco is still trying"
        → (["Carapaz","Mohoric","GallEelco is still trying"], None)  ← no possessive,
            stays as a noisy reserve; Unknown_N team gets named later by Sofia.
    """
    body = text.split(":", 1)[1] if ":" in text else text
    items = [m.strip() for m in RESERVE_ITEM_RE.findall(body)]
    items = [r for r in items if r]
    if not items:
        return [], None

    last = items[-1]
    trailing_header: Optional[Tuple[str, str]] = None

    # Case A: trailing chunk after the last reserve word, with explicit possessive.
    idx = body.rfind(last)
    trailing = body[idx + len(last):].strip() if idx >= 0 else ""
    if trailing:
        # Take the FIRST capitalized word followed by 's — closest to the reserve,
        # so we don't grab "Joe Eelco" when only "Eelco" is the player.
        m = re.search(r"\b([A-Z][a-zA-Z]+)\s*[’']s\s+(.+)$", trailing, re.DOTALL)
        if m:
            trailing_header = (m.group(1).strip(), m.group(2).strip())

    # Case B: possessive INSIDE the last reserve match (the lazy regex gobbled
    # the whole tail because there was no trailing "n)" sentinel).
    if trailing_header is None and re.search(r"[’']s\s+", last):
        m = re.match(
            r"^(.+?)\s+([A-Z][a-zA-Z]+)\s*[’']s\s+(.+)$",
            last,
            re.DOTALL,
        )
        if m:
            items[-1] = m.group(1).strip()
            trailing_header = (m.group(2).strip(), m.group(3).strip())

    return items, trailing_header


def clean_rider_name(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


# ---------------------------------------------------------------------------
# Body parser — works on a flat list of (kind, payload) events. Both the docx
# and csv parsers reduce to this same event stream so the state machine is
# shared.
# ---------------------------------------------------------------------------

def parse_events(events) -> Dict:
    """
    events: iterable of tuples
      ('header',     (player, team_name))     — explicit header paragraph
      ('table',      [rider_name, ...])       — rider table
      ('reserves',   text)                    — raw reserves paragraph
    Returns a dict { teams: [...], unresolved: [...] }.
    """
    teams: List[Dict] = []
    unresolved: List[str] = []
    pending_header: Optional[Tuple[str, str]] = None
    # Team currently waiting for its reserves paragraph:
    open_team: Optional[Dict] = None

    def make_unknown_label() -> str:
        return f"Unknown_{len(teams) + 1}"

    for kind, payload in events:
        if kind == "header":
            pending_header = payload

        elif kind == "table":
            # If the previous team never got reserves (e.g. doc ended), close it.
            if open_team is not None:
                teams.append(open_team)
                open_team = None

            if pending_header is None:
                player, name = make_unknown_label(), ""
                unresolved.append(player)
            else:
                player, name = pending_header
                pending_header = None

            open_team = {
                "player": player,
                "team_name": name,
                "riders": payload,
                "reserves": [],
                "needs_attention": player.startswith("Unknown_"),
            }

        elif kind == "reserves":
            reserves, trailing = parse_reserves_paragraph(payload)
            if open_team is not None:
                open_team["reserves"] = reserves
                teams.append(open_team)
                open_team = None
            else:
                # Reserves with no open team — record as orphan.
                unresolved.append(f"orphan_reserves:{reserves}")
            if trailing is not None:
                pending_header = trailing

    if open_team is not None:
        teams.append(open_team)

    # Backfill: when an Unknown_N team's reserves paragraph contained a
    # trailing header that we then ATTACHED TO THE NEXT team — but the
    # Unknown team itself remains nameless. That's expected: surface it.
    return {"teams": teams, "unresolved": unresolved}


# ---------------------------------------------------------------------------
# DOCX → events
# ---------------------------------------------------------------------------

def iter_docx_events(path: Path):
    doc = docx.Document(str(path))
    paras_iter = iter(doc.paragraphs)
    tables_iter = iter(doc.tables)

    # Pull year from header text.
    full = "\n".join(p.text for p in doc.paragraphs[:30])
    yield "year", find_year(full)

    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p = next(paras_iter)
            text = p.text.strip()
            if not text or is_noise(text):
                continue
            if RESERVES_LINE_RE.match(text):
                yield "reserves", text
                continue
            header = extract_header(text)
            if header:
                yield "header", header
        elif tag == "tbl":
            t = next(tables_iter)
            riders: List[str] = []
            for row in t.rows:
                first = clean_rider_name(row.cells[0].text)
                if not first or first.lower() in SKIP_HEADERS:
                    continue
                riders.append(first)
            yield "table", riders


# ---------------------------------------------------------------------------
# CSV → events
# ---------------------------------------------------------------------------

def iter_csv_events(path: Path):
    with open(path, newline="", encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))

    # Year
    year = None
    for row in rows[:5]:
        joined = " ".join(c for c in row if c).strip()
        y = find_year(joined)
        if y:
            year = y
            break
    yield "year", year

    in_table = False
    table_riders: List[str] = []

    def flush_table():
        nonlocal table_riders
        if table_riders:
            yield_table = list(table_riders)
            table_riders = []
            return ("table", yield_table)
        return None

    pending: List = []  # buffer: yields are made via this queue so we can flush_table()

    for row in rows:
        joined = " ".join(c.strip() for c in row if c.strip()).strip()
        first = (row[0] if row else "").strip()
        if not joined:
            continue
        if is_noise(joined):
            continue

        # Reserves line ends a team and may carry a trailing header.
        if RESERVES_LINE_RE.match(joined):
            flushed = flush_table()
            if flushed:
                yield flushed
            in_table = False
            yield "reserves", joined
            continue

        # Table column-headers:
        if first.lower().startswith("renner/etappe"):
            in_table = True
            continue

        # Footer rows close the rider table portion.
        if first.lower() in SKIP_HEADERS:
            flushed = flush_table()
            if flushed:
                yield flushed
            in_table = False
            continue

        # Possessive header line.
        header = extract_header(joined) or extract_header(first)
        if header and not in_table:
            flushed = flush_table()
            if flushed:
                yield flushed
            yield "header", header
            continue

        # Rider row.
        if in_table and first:
            table_riders.append(clean_rider_name(first))

    flushed = flush_table()
    if flushed:
        yield flushed


# ---------------------------------------------------------------------------
# Top-level parse
# ---------------------------------------------------------------------------

def parse_file(path: Path, fmt: str) -> Dict:
    iterator = iter_docx_events(path) if fmt == "docx" else iter_csv_events(path)
    events = list(iterator)
    year = next((p for k, p in events if k == "year"), None)
    body_events = [(k, p) for k, p in events if k != "year"]
    result = parse_events(body_events)
    return {
        "source": path.name,
        "year": year,
        "team_count": len(result["teams"]),
        "teams": result["teams"],
        "unresolved": result["unresolved"],
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

INPUTS = [
    ("docx", INPUTS_DIR / "Tour.2020 (1).docx"),
    ("csv",  INPUTS_DIR / "Tour.2021.csv"),
    ("docx", INPUTS_DIR / "Tour.2022.docx"),
    ("docx", INPUTS_DIR / "Tour.2024.docx"),
    ("docx", INPUTS_DIR / "Tour.2025.docx"),
]


def summarise(parsed: Dict) -> str:
    lines = [f"\n=== {parsed['source']} (year={parsed['year']}, teams={parsed['team_count']}) ==="]
    for t in parsed["teams"]:
        flag = " ⚠" if t.get("needs_attention") else "  "
        lines.append(
            f" {flag} {t['player']:<12} {t['team_name'][:35]:<35}  "
            f"riders={len(t['riders'])}  reserves={len(t['reserves'])}"
        )
    if parsed["unresolved"]:
        lines.append(f"  unresolved: {parsed['unresolved']}")
    return "\n".join(lines)


def main():
    out_dir = REPO / "scripts" / "validator-output"
    out_dir.mkdir(exist_ok=True)
    for fmt, path in INPUTS:
        if not path.exists():
            print(f"missing: {path}")
            continue
        parsed = parse_file(path, fmt)
        print(summarise(parsed))
        out_file = out_dir / (path.stem + ".json")
        out_file.write_text(json.dumps(parsed, indent=2, ensure_ascii=False))
    print(f"\nJSON outputs: {out_dir}")


if __name__ == "__main__":
    main()

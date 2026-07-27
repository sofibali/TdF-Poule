#!/usr/bin/env python3
"""
Fill in Tour 2026.docx tables with per-rider per-stage points from the DB.
Each of the 17 tables corresponds to one team. Fills:
  - Rider rows 1-15 (main riders): stage points when active, blank when dropped
  - Row 16 (Dagtotaal): correct team total per stage (from DB, includes reserves)
  - Row 17 (Cumulatief): running total of Dagtotaal
  - Column 21 (totaal): sum of per-rider stage points
"""
import json
import unicodedata
import shutil
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DATA = Path(__file__).parent / "teams-data.json"
SRC_DOC = Path(__file__).parent.parent / "historical-inputs/Tour 2026.docx"
OUT_DOC = Path(__file__).parent.parent / "historical-inputs/Tour 2026 filled.docx"

COMPLETED_STAGES = list(range(1, 12))   # stages 1-11 done
STAGE_COLS = list(range(1, 21))         # doc columns 1-20 cover stages 1-20
TOTAL_COL = 21                          # last col = totaal

# ── load data ─────────────────────────────────────────────────────────────────
with open(DATA) as f:
    d = json.load(f)

teams       = d["teams"]          # [{id, player_name}]
picks       = d["picks"]          # [{team_id, raw_name, rider_id, is_reserve, riders}]
dropouts    = d["dropouts"]       # [{rider_id, dropout_after_stage}]
stage_pts   = d["stagePoints"]    # [{rider_id, stage, points}]
team_sp     = d["teamStagePts"]   # [{team_id, stage, points}]

dropout_map = {row["rider_id"]: row["dropout_after_stage"] for row in dropouts}

# rider_id -> {stage -> points}
rider_pts: dict[str, dict[int, int]] = {}
for row in stage_pts:
    rid = row["rider_id"]
    s   = int(row["stage"])
    p   = int(row["points"])
    rider_pts.setdefault(rid, {})[s] = p

# team_id -> {stage -> points}
team_stage_pts: dict[str, dict[int, int]] = {}
for row in team_sp:
    tid = row["team_id"]
    s   = int(row["stage"])
    p   = int(row["points"])
    team_stage_pts.setdefault(tid, {})[s] = p

# team_id -> list of picks
team_picks: dict[str, list] = {}
for p in picks:
    team_picks.setdefault(p["team_id"], []).append(p)

# ── name matching helpers ──────────────────────────────────────────────────────
def norm(s: str) -> str:
    """Lowercase + strip accents + normalize punctuation."""
    s = (s or "").lower().replace("-", "").replace("'", "").replace("’", "").replace("'", "")
    return "".join(
        c for c in unicodedata.normalize("NFD", s)
        if unicodedata.category(c) != "Mn"
    )

# Manual overrides: normalized doc name → fragment to search for
OVERRIDES = {
    "vd poel": "van der poel",
    "vanquelin": "vauquelin",
    "uijtenbroeks": "uijtdebroeks",
    "uijtdebroeks": "uijtdebroeks",
    "healy": "healy",
    "healey": "healy",
    "mohoric": "mohoric",
    "fretin": "fretin",        # may be unmatched if not in DB
    "paret peintre": "paret",
    "fred wright": "wright",
}

def match_rider(doc_name: str, team_id: str):
    """Return the pick dict for the best-matching main rider, or None."""
    raw = norm(doc_name.strip())
    parts = raw.split()
    # strip leading single-letter initial  ("T Johannessen" → "Johannessen")
    search = " ".join(parts[1:]) if (len(parts) > 1 and len(parts[0]) <= 2) else raw
    search = OVERRIDES.get(search, search)

    main = [p for p in team_picks.get(team_id, []) if not p["is_reserve"] and p["rider_id"]]
    tokens = [t for t in search.split() if len(t) > 2]

    best, best_score = None, 0
    for p in main:
        ri = p.get("riders") or {}
        combined = norm(
            (ri.get("full_name") or "") + " "
            + (ri.get("last_name") or "") + " "
            + (p.get("raw_name") or "")
        )
        score = sum(1 for t in tokens if t in combined)
        if score > best_score:
            best_score, best = score, p

    return best if best_score > 0 else None

# ── table → team keyword mapping ──────────────────────────────────────────────
# Order matches paragraph order in the docx
TABLE_KEYWORDS = [
    "coert",    # 0
    "chiel",    # 1
    "kielen",   # 2
    "quinten",  # 3
    "lori",     # 4
    "rich",     # 5
    "hubert",   # 6
    "karin",    # 7
    "sofia",    # 8
    "gerards",  # 9
    "eelco",    # 10
    "han",      # 11
    "rein",     # 12
    "bas oud",  # 13
    "bas ot",   # 14
    "copilot",  # 15
    "claude",   # 16
]

def find_team(keyword: str):
    kn = norm(keyword)
    for t in teams:
        if kn in norm(t["player_name"]):
            return t
    # word-by-word fallback
    for word in kn.split():
        if len(word) > 2:
            for t in teams:
                if word in norm(t["player_name"]):
                    return t
    return None

# ── cell writer ───────────────────────────────────────────────────────────────
def set_cell(cell, text: str):
    """Write text into cell's first paragraph, preserving paragraph formatting."""
    para = cell.paragraphs[0]
    # clear all runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

# ── main ──────────────────────────────────────────────────────────────────────
shutil.copy(SRC_DOC, OUT_DOC)
doc = Document(OUT_DOC)

for tidx, keyword in enumerate(TABLE_KEYWORDS):
    team = find_team(keyword)
    if not team:
        print(f"[WARN] table {tidx}: no team for '{keyword}'")
        continue

    team_id = team["id"]
    print(f"\nTable {tidx}: '{keyword}' → '{team['player_name']}'")

    tbl = doc.tables[tidx]

    # ── rider rows (1-15) ──────────────────────────────────────────────────
    for row_idx in range(1, 16):
        doc_name = tbl.rows[row_idx].cells[0].text.strip()
        pick = match_rider(doc_name, team_id)

        if not pick:
            print(f"  [WARN] row {row_idx}: no match for '{doc_name}'")
            continue

        rider_id    = pick["rider_id"]
        drop_stage  = dropout_map.get(rider_id)   # dropout_after_stage or None
        rider_total = 0

        for col_idx, stage in enumerate(STAGE_COLS, start=1):
            if stage not in COMPLETED_STAGES:
                break   # leave future stages blank

            # rider inactive if dropped before this stage
            if drop_stage is not None and drop_stage < stage:
                continue   # leave blank

            pts = rider_pts.get(rider_id, {}).get(stage, 0)
            if pts:
                set_cell(tbl.rows[row_idx].cells[col_idx], str(pts))
                rider_total += pts

        if rider_total:
            set_cell(tbl.rows[row_idx].cells[TOTAL_COL], str(rider_total))

    # ── Dagtotaal (row 16) and Cumulatief (row 17) ──────────────────────────
    cumulative = 0
    team_total = 0
    for col_idx, stage in enumerate(STAGE_COLS, start=1):
        dagtotaal = team_stage_pts.get(team_id, {}).get(stage, 0)
        if stage in COMPLETED_STAGES and dagtotaal:
            set_cell(tbl.rows[16].cells[col_idx], str(dagtotaal))
            cumulative += dagtotaal
            set_cell(tbl.rows[17].cells[col_idx], str(cumulative))
        # else leave blank

    team_total = sum(
        team_stage_pts.get(team_id, {}).get(s, 0) for s in COMPLETED_STAGES
    )
    if team_total:
        set_cell(tbl.rows[16].cells[TOTAL_COL], str(team_total))
        set_cell(tbl.rows[17].cells[TOTAL_COL], str(team_total))

doc.save(OUT_DOC)
print(f"\nSaved → {OUT_DOC}")
